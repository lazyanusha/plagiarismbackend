import os
import re
import sys
import logging
import requests
from io import BytesIO
import numpy as np
import nltk
from nltk.tokenize import sent_tokenize
from sentence_transformers import SentenceTransformer
from itertools import groupby
from operator import itemgetter
from docx import Document
import PyPDF2
from app.algorithm.citation_checker import classify_citation_status  # your import

# Setup logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

# Setup nltk
nltk.download('punkt', quiet=True)

EPSILON = sys.float_info.epsilon
MODEL_NAME = 'all-MiniLM-L6-v2'
model = SentenceTransformer(MODEL_NAME)

# -----------------------------
# File Reader Functions
# -----------------------------

def merge_broken_lines(lines):
    merged_lines = []
    buffer = ""
    for line in lines:
        stripped = line.strip()
        ends_sentence = re.search(r'[.!?]["\']?$', stripped)
        if not buffer:
            buffer = stripped
        else:
            if not ends_sentence:
                buffer += " " + stripped
            else:
                buffer += " " + stripped
                merged_lines.append(buffer.strip())
                buffer = ""
    if buffer:
        merged_lines.append(buffer.strip())
    return merged_lines

def split_into_sentences(text):
    return sent_tokenize(text)

def read_txt_from_string(text):
    lines = text.split('\n')
    return merge_broken_lines(lines)

def read_docx_from_bytes(bytes_data):
    doc = Document(BytesIO(bytes_data))
    lines = [para.text for para in doc.paragraphs if para.text.strip()]
    return merge_broken_lines(lines)

def read_pdf_from_bytes(bytes_data):
    text = ''
    reader = PyPDF2.PdfReader(BytesIO(bytes_data))
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + '\n'
    lines = text.split('\n')
    return merge_broken_lines(lines)

def read_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as f:
        text = f.read()
    lines = text.split('\n')
    merged = merge_broken_lines(lines)
    return merged

def read_docx(file_path):
    doc = Document(file_path)
    lines = [para.text for para in doc.paragraphs if para.text.strip()]
    merged = merge_broken_lines(lines)
    return merged

def read_pdf(file_path):
    text = ''
    with open(file_path, 'rb') as f:
        reader = PyPDF2.PdfReader(f)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + '\n'
    lines = text.split('\n')
    merged = merge_broken_lines(lines)
    return merged

def read_raw_lines(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.txt':
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    elif ext == '.docx':
        doc = Document(file_path)
        lines = [para.text for para in doc.paragraphs]
    elif ext == '.pdf':
        lines = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    lines.extend(page_text.split('\n'))
    else:
        with open(file_path, 'r', encoding='utf-8') as f:
            lines = f.readlines()
    return lines

def read_file(input_source):
    """
    input_source can be:
    - local file path (str)
    - URL (str)
    """
    try:
        if input_source.startswith('http://') or input_source.startswith('https://'):
            logging.info(f"Fetching file from URL: {input_source}")
            response = requests.get(input_source)
            response.raise_for_status()
            content_type = response.headers.get('Content-Type', '').lower()
            content = response.content
            if 'application/pdf' in content_type or input_source.lower().endswith('.pdf'):
                logging.info("Detected PDF content")
                merged = read_pdf_from_bytes(content)
            elif ('wordprocessingml.document' in content_type) or input_source.lower().endswith('.docx'):
                logging.info("Detected DOCX content")
                merged = read_docx_from_bytes(content)
            else:  # treat as text
                logging.info("Treating content as plain text")
                text = response.text
                merged = read_txt_from_string(text)
        elif os.path.isfile(input_source):
            ext = os.path.splitext(input_source)[1].lower()
            logging.info(f"Reading local file: {input_source} with extension {ext}")
            if ext == '.txt':
                merged = read_txt(input_source)
            elif ext == '.docx':
                merged = read_docx(input_source)
            elif ext == '.pdf':
                merged = read_pdf(input_source)
            else:
                logging.warning(f"Unsupported extension {ext}, falling back to reading as text")
                merged = read_txt(input_source)
        else:
            logging.error(f"Input source is not a valid file path or URL: {input_source}")
            merged = []
    except Exception as e:
        logging.error(f"Error reading input source {input_source}: {e}")
        merged = []

    sentences = []
    for line in merged:
        if len(line.split()) > 10:
            sentences.extend(split_into_sentences(line))
        else:
            sentences.append(line)
    return sentences

# -----------------------------
# Similarity & Detection Logic
# -----------------------------

def get_sentence_embeddings(sentences, model):
    return model.encode(sentences, convert_to_numpy=True)

def cosine_similarity(matrix1, matrix2):
    matrix1_norm = np.linalg.norm(matrix1, axis=1, keepdims=True) + EPSILON
    matrix2_norm = np.linalg.norm(matrix2, axis=1, keepdims=True) + EPSILON
    dot_product = np.dot(matrix1, matrix2.T)
    norm_product = np.dot(matrix1_norm, matrix2_norm.T)
    return dot_product / norm_product

def compute_similarity_matrix(embeddings1, embeddings2):
    return cosine_similarity(embeddings1, embeddings2)

def extract_plagiarized_pairs(sentences1, sentences2, similarity_matrix, threshold=0.8):
    exact_threshold = 0.95
    plagiarized_pairs = []
    for i, row in enumerate(similarity_matrix):
        j = np.argmax(row)
        similarity_score = row[j]
        if similarity_score >= threshold:
            match_type = "exact" if similarity_score >= exact_threshold else "paraphrased"
            plagiarized_pairs.append({
                "doc1_idx": i,
                "doc1_sentence": sentences1[i],
                "doc2_idx": j,
                "doc2_sentence": sentences2[j],
                "similarity": similarity_score,
                "type": match_type
            })
    return plagiarized_pairs

def group_consecutive_indices(indices):
    groups = []
    for k, g in groupby(enumerate(indices), lambda ix: ix[0] - ix[1]):
        group = list(map(itemgetter(1), g))
        groups.append(group)
    return groups

def get_plagiarism_blocks(matched_pairs, sentences_doc1, sentences_doc2):
    matched_pairs.sort(key=lambda x: x['doc1_idx'])
    plag_indices = [p['doc1_idx'] for p in matched_pairs]
    grouped_indices = group_consecutive_indices(plag_indices)
    blocks = []
    for group in grouped_indices:
        doc1_block_sentences = [sentences_doc1[i] for i in group]
        doc1_block_text = " ".join(doc1_block_sentences)
        group_pairs = [p for p in matched_pairs if p['doc1_idx'] in group]
        doc2_indices = sorted(set(p['doc2_idx'] for p in group_pairs))
        doc2_block_sentences = [sentences_doc2[j] for j in doc2_indices]
        doc2_block_text = " ".join(doc2_block_sentences)
        avg_similarity = sum(p['similarity'] for p in group_pairs) / len(group_pairs)
        blocks.append({
            'doc1_text': doc1_block_text,
            'doc2_text': doc2_block_text,
            'avg_similarity': avg_similarity,
            'doc1_indices': group,
            'doc2_indices': doc2_indices,
        })
    return blocks

# -----------------------------
# Main plagiarism detection with citation checking
# -----------------------------

def get_plagiarism_report(file_path1, file_path2, threshold=0.8, display_name=None):
    logging.info(f"Generating plagiarism report for '{file_path1}' vs '{file_path2}'")
    doc1 = read_file(file_path1)
    doc2 = read_file(file_path2)
    user_basename = os.path.basename(file_path1)
    if not doc1:
        logging.warning(f"No text extracted from {file_path1}")
    if not doc2:
        logging.warning(f"No text extracted from {file_path2}")

    embeddings_doc1 = get_sentence_embeddings(doc1, model)
    embeddings_doc2 = get_sentence_embeddings(doc2, model)

    similarity_matrix = compute_similarity_matrix(embeddings_doc1, embeddings_doc2)

    exact_threshold = 0.95
    num_sentences = len(doc1)
    if num_sentences == 0:
        logging.warning("No sentences found in first document; returning empty report")
        return {
            "uploaded_filename": user_basename, 
            "filename": display_name or os.path.basename(file_path2),
            "exact_score": 0.0,
            "partial_score": 0.0,
            "unique_score": 1.0,
            "total_score": 1.0,
            "exact_matches": [],
            "partial_matches": [],
            "matched_pairs": []
        }

    exact_matches = []
    partial_matches = []
    matched_pairs = []
    unique_count = 0

    for i, row in enumerate(similarity_matrix):
        max_j = np.argmax(row)
        max_sim = row[max_j]
        if max_sim >= exact_threshold:
            exact_matches.append(doc1[i])
            matched_pairs.append({
                "doc1_idx": i,
                "doc1_sentence": doc1[i],
                "doc2_idx": max_j,
                "doc2_sentence": doc2[max_j],
                "similarity": max_sim,
                "type": "exact",
                "source_file": display_name or os.path.basename(file_path2)
            })
        elif max_sim >= threshold:
            partial_matches.append(doc1[i])
            matched_pairs.append({
                "doc1_idx": i,
                "doc1_sentence": doc1[i],
                "doc2_idx": max_j,
                "doc2_sentence": doc2[max_j],
                "similarity": max_sim,
                "type": "partial",
                "source_file": display_name or os.path.basename(file_path2)
            })
        else:
            unique_count += 1

    try:
        doc2_lines = read_raw_lines(file_path2)
    except Exception as e:
        logging.warning(f"Could not read raw lines from {file_path2} for citation checking: {e}")
        doc2_lines = []

    matched_pairs = classify_citation_status(matched_pairs, doc1, doc2, doc2_lines)

    x = len(exact_matches) / num_sentences
    y = len(partial_matches) / num_sentences
    z = unique_count / num_sentences

    report = {
        "filename": display_name or os.path.basename(file_path2),
        "exact_score": round(x, 4),
        "partial_score": round(y, 4),
        "unique_score": round(z, 4),
        "total_score": 1.0,
        "exact_matches": exact_matches,
        "partial_matches": partial_matches,
        "matched_pairs": matched_pairs
    }

    logging.info(f"Plagiarism report generated for '{display_name or file_path2}'")
    return report


# __main__ block for quick local testing
if __name__ == '__main__':
    file1 = 'Proposal.docx'
    file2 = 'Proposal.docx'
    from pprint import pprint
    report = get_plagiarism_report(file1, file2)
    pprint(report)
   
